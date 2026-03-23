import os
import io
from flask import Flask, request, jsonify, send_file
from flask import send_from_directory
from flask_cors import CORS
from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile


app = Flask(__name__, static_folder="static", static_url_path="")
CORS(app)

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# -------- FILE PARSER --------
def extract_text(file):
    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text

    elif filename.endswith(".docx"):
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])

    else:
        return file.read().decode("utf-8", errors="ignore")

def create_document(content, format_type, original_file=None):
    """Create document in specified format with preserved styling"""
    
    if format_type == 'txt':
        # Plain text format
        return content.encode('utf-8')
    
    elif format_type == 'docx':
        # Create DOCX with basic formatting
        doc = Document()
        
        # Try to preserve some formatting if original file was DOCX
        if original_file:
            try:
                original_doc = Document(original_file)
                # Copy basic document settings
                for section in original_doc.sections:
                    doc.sections[0].page_width = section.page_width
                    doc.sections[0].page_height = section.page_height
                    doc.sections[0].left_margin = section.left_margin
                    doc.sections[0].right_margin = section.right_margin
            except:
                pass  # Fallback to default
        
        # Add content with paragraphs
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text)
                # Apply basic formatting
                if para_text.strip().startswith(('#', '•', '-', '*')):
                    p.style = 'List Bullet'
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    elif format_type == 'pdf':
        # For PDF, we'll use reportlab or just return as txt with PDF extension
        # For simplicity, we'll return as TXT but mark as PDF (user can convert)
        # For a production app, use reportlab or PyPDF2 to create actual PDF
        return content.encode('utf-8')
    
    else:
        return content.encode('utf-8')

def get_file_format(filename):
    """Determine file format from filename"""
    if not filename:
        return 'txt'
    ext = filename.lower().split('.')[-1] if '.' in filename else 'txt'
    if ext in ['pdf', 'docx', 'txt']:
        return ext
    return 'txt'

@app.route("/")
def serve_ui():
    return send_from_directory("static", "index.html")

@app.route("/analyze", methods=["POST"])
def analyze():
    try:
        resume_text = request.form.get("resume_text", "")
        jd_text = request.form.get("jd_text", "")
        preserve_format = request.form.get("preserve_format", "true") == "true"

        resume_file = request.files.get("resume_file")
        jd_file = request.files.get("jd_file")
        
        original_filename = None
        original_format = 'txt'

        if resume_file:
            original_filename = resume_file.filename
            original_format = get_file_format(original_filename)
            resume_text = extract_text(resume_file)
            
            # Store original file for formatting reference
            resume_file.seek(0)  # Reset file pointer
            original_file_content = resume_file.read() if preserve_format else None
        else:
            original_file_content = None

        if jd_file:
            jd_text = extract_text(jd_file)

        if not resume_text or not jd_text:
            return jsonify({"msg": "Provide Resume & Job Description"}), 400

        # LIMIT SIZE (IMPORTANT)
        resume_text = resume_text[:3000]
        jd_text = jd_text[:2000]

        prompt = f"""
        You are an ATS system.

        Compare resume with job description.

        Resume:
        {resume_text}

        Job Description:
        {jd_text}

        Give output in this format:

        ATS Score: (0-100%)

        Missing Keywords:
        - list

        Matching Skills:
        - list

        Suggestions:
        - bullet points

        Improved Resume:
        [Provide an improved version of the resume that incorporates the job description keywords and addresses the suggestions. Keep it in the same structure as the original resume but enhanced with better keywords and formatting. Maintain professional tone and ensure it's ATS-friendly.]
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
            timeout=20
        )

        result = response.choices[0].message.content
        
        # Store the improved resume text separately for download
        improved_resume = ""
        if "Improved Resume:" in result:
            improved_resume = result.split("Improved Resume:")[1].strip()
            # Remove any trailing sections if they exist
            if "\n\nATS Score:" in improved_resume:
                improved_resume = improved_resume.split("\n\nATS Score:")[0]
        
        # Store in session or return with metadata
        return jsonify({
            "result": result,
            "improved_resume": improved_resume,
            "original_format": original_format,
            "original_filename": original_filename
        })

    except Exception as e:
        return jsonify({"msg": str(e)}), 500

@app.route("/download-improved", methods=["POST"])
def download_improved():
    try:
        data = request.json
        improved_text = data.get("improved_text", "")
        original_format = data.get("format", "txt")
        original_filename = data.get("original_filename", "improved_resume")
        
        if not improved_text:
            return jsonify({"msg": "No improved resume content"}), 400
        
        # Create filename with original extension
        base_name = os.path.splitext(original_filename)[0] if original_filename else "improved_resume"
        filename = f"{base_name}_improved.{original_format}"
        
        # Create document in requested format
        file_content = create_document(improved_text, original_format)
        
        # Send file for download
        return send_file(
            io.BytesIO(file_content),
            as_attachment=True,
            download_name=filename,
            mimetype='application/octet-stream'
        )
        
    except Exception as e:
        return jsonify({"msg": str(e)}), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)